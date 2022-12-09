import datetime
import os
import re
import sys
import time

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from lxml import etree

from exam import exam
from login import chaoxing_login


class Chaoxing_spider():

    def __init__(self):
        self.session = chaoxing_login().login()
        self.my_teach_headers = {
            'Host': 'mooc1-1.chaoxing.com',
            # 'Referer': 'http://mooc1-1.chaoxing.com/visit/interaction?s=e9059bca0eca12ef882b78f6a497cdc9',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest',
        }
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
            'Host': 'mooc1-1.chaoxing.com',
        }
        self.class_dic = {}
        self.statistics_info_dic = {}
        self.schoolid = '18638'  # 自己学校Id，学校不同Id不同

    # 获取班级链接及名字，并存入一个字典
    def get_class_url_name(self):
        params = {
            'isAjax': 'true'
        }
        url = 'http://mooc1-1.chaoxing.com/visit/courses/teach'
        self.session.headers = self.my_teach_headers
        response = self.session.get(url=url, params=params)
        # print(response.text)
        text = response.text
        href_name = re.findall('<a class="courseName" href=(.*?) target="_blank" title=(.*?) >', text)
        for i in href_name:
            href, name = i
            href = 'https://mooc1-1.chaoxing.com' + href.replace("'", "").replace('"', '')
            name = name.replace("'", "").replace('"', '')
            self.class_dic[name] = href
        return ''

    # 获取成绩统计的信息，并获取信息的题头
    def get_statistics_info(self):
        title_url = 'https://mooc1-1.chaoxing.com/moocAnalysis/analysisScore?classId={}&courseId={}&ut=t&cpi={}&openc={}'.format(
            self.classId, self.courseId, self.cpi, self.openc)
        r = self.session.get(url=title_url)
        text = r.text
        html = etree.HTML(text)
        th_list = html.xpath('//tr[@id="commonthead"]/th')
        title_list = []
        for th in th_list:
            i = th.xpath('span/text()')
            if len(i) == 1:
                title_list.append(i[0])
            else:
                title_list.append(i[0] + i[1])
        # 创建字典
        for i in title_list:
            self.statistics_info_dic[i] = []
        # print(self.statistics_info_dic)
        data = {
            'courseId': self.courseId,
            'classId': self.classId,
            'pageSize': '30',
            'sw': '',
            'pageNum': '1',
            'fid': '0',
            'sortType': '',
            'order': '',
            'test': '0',
            'isSimple': '0',
            'openc': self.openc,
        }
        self.session.headers = {
            'Host': 'mooc1-1.chaoxing.com',
            'Origin': 'https://mooc1-1.chaoxing.com',
            'Referer': self.statistics_url,
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest',
        }
        url = 'https://mooc1-1.chaoxing.com/moocAnalysis/analysisScoreData'
        p = 1
        while True:
            data['pageNum'] = str(p)
            p += 1
            response = self.session.post(url=url, data=data)
            # print(response.text)
            text = response.text
            if text == '':
                break
            else:
                self.statistics_info_handle(text, title_list)
        # 保存文件
        df = pd.DataFrame(self.statistics_info_dic)
        print(df)
        file_path = r'data\statistics\{}.xlsx'.format(self.class_name + str(int(time.time() * 1000)))
        b = os.path.exists(file_path)
        if b == False:
            print('{}excel文件正在保存......'.format(self.class_name))
            df.to_excel(file_path)
            print('{}excel文件保存成功.'.format(self.class_name))
        else:
            print('路径为：{}\n'
                  'excel文件存在，无法保存\n'
                  '程序将自动退出'.format(file_path))
            sys.exit()

    # 根据信息标头对每个人信息进行提取
    def statistics_info_handle(self, text, title_list):
        # print(title_list)
        html = etree.HTML(text)
        # 把每次提取到的成绩存入info_lint
        info_list = []
        for i in range(1, len(title_list) + 1):
            if i == 1:
                info = html.xpath('//tr/td[1]/span/@title')
                # print(info)
            else:
                info = html.xpath('//tr/td[{}]/span/text()'.format(i))
                for j in range(len(info)):
                    info[j] = info[j].replace('\t', '').replace('\n', '').replace('\r', '').strip()
                # print(info)
            info_list.append(info)
            # print(info)
        # print(info_list)
        dic = dict(zip(title_list, info_list))
        for i in title_list:
            self.statistics_info_dic[i] = self.statistics_info_dic[i] + dic[i]
        # print(self.statistics_info_dic)

    # 模板试卷
    def template_exam(self):
        base_url = 'https://mooc1-1.chaoxing.com/exam/loadPaperTemplate'
        params = {
            'courseId': self.courseId,
            'start': '0',
            'examsystem': '0',
            'isCustomPaper': 'false',
            'cpi': self.cpi,
            'openc': self.openc,
            'qbanksystem': '0',
            'qbankbackurl': '',
        }
        response = self.session.get(url=base_url, params=params)
        text = response.text
        html = etree.HTML(text)
        tr_list = html.xpath('//tbody[@id="tableId"]/tr')
        name_url_list = []
        n = 0
        for tr in tr_list:
            name = tr.xpath('td[1]/text()')[0]
            url_info = tr.xpath('td[5]/a/@onclick')[0]
            template_url = 'https://mooc1-1.chaoxing.com' + re.findall('"(.*?)"', url_info)[0]
            name_url_list.append((name, template_url))
            print('序号:{}  名字:{}'.format(n, name))
            n += 1
        print('输入负数将退出!')
        while True:
            num = int(input("请输入你要选择模板的序号:"))
            if num < len(name_url_list) and num >= 0:
                template_url = name_url_list[num][1]
                break
            elif num >= len(name_url_list):
                print('输入数字有误!')
            else:
                sys.exit()
        paper_num = int(input("请输入要组成试卷数量:"))
        exam().template_exam(template_url, paper_num)

    # 试卷库
    def exam_library(self):
        url = 'https://mooc1-1.chaoxing.com/exam/reVerSionPaperList'
        params = {
            'courseId': self.courseId,
            'classId': self.classId,
            'ut': 't',
            'examsystem': '0',
            'cpi': self.cpi,
            'openc': self.openc,
        }
        response = self.session.get(url=url, params=params)
        text = response.text
        href_name_list = re.findall('<a class="title" href="(.*?)">(.*?)</a>', text)
        for i in href_name_list:
            href, name = i
            url = '{}{}'.format('https://mooc1-1.chaoxing.com', href)
            content = self.session.get(url).content
            html = etree.HTML(content)
            url = 'https://mooc1-1.chaoxing.com' + html.xpath('//*[@id="RightCon"]/div[2]/ul/li/div[2]/p/a/@href')[0]
            self.download_paper(url, name)

    def download_paper(self, url, name):
        paper_dict = {}
        print('-' * 20)
        response = self.session.get(url)
        text = response.text
        # 每个答题信息
        list = re.findall('<div class="Cy_TItle1">(.*?)<!--  目录 结束-->', text, re.S)
        for i in list:
            html = etree.HTML(i)
            # 题目大标题
            subject_title = html.xpath('//h2/text()|//h2/em/text()')
            # print(subject_title)
            # 把题目大标题变为元组
            subject_list_tuple = tuple(subject_title)
            # 大标题作为第一层key
            paper_dict[subject_list_tuple] = {}
            # 题目列表
            subject_list = html.xpath('//div[@class="TiMu"]/div[@name="certainTitle"]')
            for subject in subject_list:
                # 题目信息
                subject_detailed = subject.xpath(
                    'div[1]/i/text()|div[1]/div/text()|div[1]/div/p/text()|div[1]/div/img/@src')
                # print(subject_detailed)
                # 把题目信息添加
                subject_detailed_tuple = tuple(subject_detailed)
                paper_dict[subject_list_tuple][subject_detailed_tuple] = {}
                # 选项
                option_list = subject.xpath('ul/li')
                paper_dict[subject_list_tuple][subject_detailed_tuple]['选项'] = []
                if option_list != []:
                    for option in option_list:
                        option1 = option.xpath('i/text()')
                        option1_content_list = option.xpath('div/a/text()|div/a/img/@src|div/a/p/text()')
                        option1.extend(option1_content_list)
                        paper_dict[subject_list_tuple][subject_detailed_tuple]['选项'].append(option1)
                        # print(option1)
                # 答案
                answer_list = subject.xpath(
                    'div[2]/div[1]/span/div/text()|div[2]/div[1]/span/div/img/@src|div[2]/span/text()|div[2]/div[1]/span/div/p/img/@src')
                if answer_list != []:
                    if len(answer_list) == 1:
                        answer_list = answer_list[0].replace('正确答案：', '').strip()
                        # print('答案', answer_list)
                    else:
                        answer_list1 = []
                        for i in range(len(answer_list)):
                            m = answer_list[i].strip()
                            if m != '':
                                answer_list1.append(m)

                        answer_list = answer_list1
                        # print('答案', answer_list)
                    paper_dict[subject_list_tuple][subject_detailed_tuple]['答案或分析'] = answer_list
                else:
                    analysis_list = subject.xpath('div[3]/span/img/@src')
                    # print('分析', analysis_list)
                    paper_dict[subject_list_tuple][subject_detailed_tuple]['答案或分析'] = analysis_list
        time_stamp = int(time.time())
        year = datetime.datetime.now().year
        month = datetime.datetime.now().month

        def paper():
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36', }
            self.session.headers = headers
            # 试卷document
            document = Document(r'core\template.docx')
            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
            if month <= 7:
                title0 = '浙江大学宁波理工学院{}–{}学年第二学期'.format(year, year + 1)
            else:
                title0 = '浙江大学宁波理工学院{}–{}学年第一学期'.format(year, year + 1)
            paragraph0 = document.add_paragraph('')
            run = paragraph0.add_run(title0)
            run.font.size = Pt(18)  # 18为小二
            run.bold = True  # 字体加黑
            # 设置行距
            paragraph_format = paragraph0.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            # 居中
            paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title1 = '《{}》课程期末考试试卷 (A或B)'.format(self.class_name)
            paragraph1 = document.add_paragraph('')
            run1 = paragraph1.add_run(title1)
            run1.font.size = Pt(15)  # 15为小三
            run1.bold = True  # 字体加黑
            # 设置行距
            paragraph_format = paragraph1.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            # 居中
            paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph2 = document.add_paragraph('')
            run2 = paragraph2.add_run("开课单位：____________，考试形式：闭（开、半开)卷，允许带_________入场\n"
                                      "考试日期：____________年____月____日，考试所需时间：_______分钟\n"
                                      "考生姓名：_______学号：_______考生所在学院（系）：______专业班级______")
            run2.line_spacing = 1.5  # 1.5倍行距
            run2.font.size = Pt(12)  # 12为小四
            # 设置行距
            paragraph_format = paragraph2.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            # 表格
            table = document.add_table(rows=4, cols=len(paper_dict) + 2, style='Table Grid')
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 表头
            table.rows[2].cells[0].text = '得分'
            table.rows[3].cells[0].text = '评卷人'
            hc0 = table.rows[0].cells
            hc1 = table.rows[1].cells
            hc0[0].text = '题序'
            hc1[0].text = '题型'
            n = 1
            for key, _ in paper_dict.items():
                texparagraph0, text1 = key[0].split('、')
                hc0[n].text = texparagraph0
                hc1[n].text = text1
                n += 1
            hc0[n].text = '总分'
            for m in range(4):
                # for n in range(len(paper_dict)+2):
                table.rows[m].height = Cm(1)
            # 合并单元格
            table.cell(1, len(paper_dict) + 1).merge(table.cell(3, len(paper_dict) + 1))

            # 第一层key代表大题目标题
            for key, value in paper_dict.items():
                document.add_paragraph(key[0] + key[1])
                # print(key)
                # 第二层key为大题中每个小题的题干信息
                for key, value in value.items():
                    # print(key)
                    # 题干
                    '''开启一个新段落'''
                    paragraph = document.add_paragraph('')
                    run = document.paragraphs[-1].add_run()
                    # 设置行距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    for i in key:
                        if '.png' in i:
                            response = self.session.get(i)
                            content = response.content
                            with open('core\paper.png', 'wb') as f:
                                f.write(content)
                            run.add_picture('core\paper.png')
                        else:
                            run.add_text(i)
                    # print(value)
                    # 选项
                    choice_list = value['选项']
                    if choice_list != []:
                        for choice in choice_list:
                            '''开启一个新段落'''
                            paragraph = document.add_paragraph('')
                            run = document.paragraphs[-1].add_run()
                            # 设置行距
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.line_spacing = 1.5  # 1.5倍行距
                            for i in choice:
                                if '.png' in i:
                                    response = self.session.get(i)
                                    content = response.content
                                    with open('core\paper.png', 'wb') as f:
                                        f.write(content)
                                    run.add_picture('core\paper.png')
                                else:
                                    run.add_text(i)
            document.save('data/paper/{}_{}.docx'.format(name, time_stamp))

        def answer():
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36', }
            self.session.headers = headers
            # 答案document
            document = Document()
            for key, value in paper_dict.items():
                document.add_heading(key[0], 1)
                for key, value in value.items():
                    # 开启一个新段落
                    document.add_paragraph('')
                    run = document.paragraphs[-1].add_run()
                    run.add_text('{},'.format(key[0]))
                    answer_list1 = value['答案或分析']
                    for i in answer_list1:
                        if '.png' in i:
                            response = self.session.get(i)
                            content = response.content
                            with open('core\paper.png', 'wb') as f:
                                f.write(content)
                            run.add_picture('core\paper.png')
                        else:
                            run.add_text(i)
            document.save('data/paper/{}_{}答案.docx'.format(name, time_stamp))

        print('试卷 {} 正在下载'.format(name))
        paper()
        print('试卷 {} 下载完成'.format(name))
        print('试卷 {} 答案正在下载'.format(name))
        answer()
        print('试卷 {} 答案下载完成'.format(name))

    # 试卷随机组成
    def random_exam(self, url):
        exam().random_exam(url)

    def mkdir(self):
        b = os.path.exists('data')
        if b == False:
            os.makedirs('data')
        b = os.path.exists('data/paper')
        if b == False:
            os.makedirs('data/paper')
        b = os.path.exists('data/statistics')
        if b == False:
            os.makedirs('data/statistics')

    # 课程选择及获取courseId, classId, cpi, openc四个参数
    def choice_class(self):
        self.mkdir()
        self.get_class_url_name()
        n = 0
        print('-' * 20)
        print('序号     课程')
        for name in self.class_dic.keys():
            print('{}       {}'.format(n, name))
            n += 1
        print('-' * 20)
        class_name_list = [class_name for class_name in self.class_dic.keys()]
        size = len(class_name_list)
        num = int(input("如果输入负数程序将退出\n请输入你要获取成绩的课程的序号:"))
        if num >= 0:
            while True:
                if num < size:
                    break
                else:
                    print('警告===》输入序号有误请重新选择')
                    num = int(input("请重新输入你要获取成绩的课程的序号:"))
            class_name = class_name_list[num]
            self.class_name = class_name
            url = self.class_dic[class_name]
            self.session.headers = self.headers
            response = self.session.get(url=url, allow_redirects=True)
            text = response.text
            # 统计成绩链接
            base_statistics_url = re.findall('<a.*?href="(.*?)".*?title="统计">', text)
            if base_statistics_url == []:
                print('您没有权限对==》{}《==课程操作'.format(class_name))
                print('程序将自我退出.......')
                print('程序退出成功......')
                sys.exit()
            else:
                self.statistics_url = 'https://mooc1-1.chaoxing.com' + base_statistics_url[0]
                self.courseId, self.classId, self.cpi, self.openc = \
                    re.findall(
                        r'.*?courseId=(\w*[0-9]\w*)&classId=(\w*[0-9]\w*)&.*?cpi=(\w*[0-9]\w*)&openc=(\w*[0-9]\w*)',
                        self.statistics_url)[0]

                def choice():
                    num = input('0代表成绩爬取\n'
                                '1代表下载试卷\n'
                                '2代表试卷的模板组卷\n'
                                '3代表自定义组卷\n'
                                '其它字符代表退出\n'
                                '  请输入数字:')
                    if num == '0':
                        self.get_statistics_info()
                        print('任务完成！\n'
                              '将返回主界面。')
                    elif num == '1':
                        self.exam_library()
                        print('任务完成！\n'
                              '将返回主界面。')
                    elif num == '2':
                        self.template_exam()
                        print('任务完成！\n'
                              '将返回主界面。')
                    elif num == '3':
                        print("正在运行，请稍等！")
                        self.random_exam(url)
                        print('任务完成！\n'
                              '将返回主界面。')
                    else:
                        sys.exit()

                while True:
                    choice()
                    print('*'*20)
        else:
            print('程序将要退出......')
            print('程序退出成功......')
            sys.exit()


cx_spider = Chaoxing_spider()
cx_spider.choice_class()

# class JW_spider():
#
#     def __init__(self):
#         self.session = Educational_administration_system_login().login()


# pyinstaller -F spiders.py